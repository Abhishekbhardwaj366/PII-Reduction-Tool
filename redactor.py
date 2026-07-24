import re
from docx import Document

PII_PATTERNS = {
    "email": re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"),
    "phone": re.compile(r"\b(?:\+?\d{1,3}[-\s]?)?(?:\d{3,5}[-\s]?\d{3,5}[-\s]?\d{3,5})\b"),
    "ssn": re.compile(r"\b\d{3}-?\d{2}-?\d{4}\b"),
    "credit_card": re.compile(r"\b(?:\d[ -]*?){13,19}\b"),
    "ip_address": re.compile(r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b"),
    "dob": re.compile(r"\b(?:\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}|\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2})\b"),
    "full_name": re.compile(r"\b([A-Z][a-z]+)\s+([A-Z][a-z]+)\b"),
    "company": re.compile(r"\b([A-Z][A-Za-z0-9& ]+?)\s+(?:Ltd|Limited|LLC|Inc|Corporation|Corp|Pvt\.?\s+Ltd)\b"),
    "address": re.compile(r"\b\d{1,4}\s+[A-Za-z0-9 .,-]+(?:Road|Rd|Street|St|Nagar|Marg|Colony|Avenue|Ave)\b"),
}

FAKE_VALUES = {
    "email": "john.doe@example.com",
    "phone": "+91 9876543210",
    "ssn": "000-00-0000",
    "credit_card": "4111 1111 1111 1111",
    "ip_address": "192.0.2.1",
    "dob": "01-01-1990",
    "full_name": "John Doe",
    "company": "Acme Corp",
    "address": "123 Example Road",
}

def redact_match(match, pii_type):
    return FAKE_VALUES[pii_type]

def redact_text(text):
    redacted = text
    pii_hits = []
    for pii_type, pattern in PII_PATTERNS.items():
        def _replace(m):
            pii_hits.append((pii_type, m.group(0)))
            return redact_match(m, pii_type)
        redacted = pattern.sub(_replace, redacted)
    return redacted, pii_hits

def read_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def write_docx(text, path="redacted_output.docx"):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(path)

def write_hits_csv(hits, path="pii_hits.csv"):
    with open(path, "w", encoding="utf-8") as f:
        f.write("pii_type,original\n")
        for pii_type, original in hits:
            f.write(f"{pii_type},{original.replace(chr(10), ' ')}\n")

def main():
    original_text = read_docx(r"C:\Users\abhi2\Downloads\Red Herring Prospectus.docx")
    redacted_text, hits = redact_text(original_text)
    write_hits_csv(hits, "pii_hits.csv")
    write_docx(redacted_text, "redacted_output.docx")
    print("Redaction completed.")

if __name__ == "__main__":
    main()